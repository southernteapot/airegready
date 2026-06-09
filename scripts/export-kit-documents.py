#!/usr/bin/env python3
"""Convert kit markdown documents to DOCX matching the existing deliverable style.

Usage: python3 scripts/export-kit-documents.py JOBFILE

JOBFILE is a Python literal list of jobs:
  {"md": path, "docx": path, "template": path-to-existing-kit-docx, "footer": text}

PDF conversion is done separately with LibreOffice (soffice --headless).
"""
import re
import shutil
import sys

from docx import Document
from docx.oxml import parse_xml
from docx.oxml.ns import nsmap, qn
from docx.shared import Pt

NS = 'xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"'

TBL_PR = (
    '<w:tblPr %s>'
    '<w:tblW w:type="auto" w:w="0"/>'
    '<w:jc w:val="center"/>'
    '<w:tblLayout w:type="autofit"/>'
    '<w:tblLook w:firstColumn="1" w:firstRow="1" w:lastColumn="0" w:lastRow="0" w:noHBand="0" w:noVBand="1" w:val="04A0"/>'
    '<w:tblBorders>'
    '<w:top w:val="single" w:sz="4" w:space="0" w:color="D7E5F8"/>'
    '<w:left w:val="single" w:sz="4" w:space="0" w:color="D7E5F8"/>'
    '<w:bottom w:val="single" w:sz="4" w:space="0" w:color="D7E5F8"/>'
    '<w:right w:val="single" w:sz="4" w:space="0" w:color="D7E5F8"/>'
    '<w:insideH w:val="single" w:sz="4" w:space="0" w:color="D7E5F8"/>'
    '<w:insideV w:val="single" w:sz="4" w:space="0" w:color="D7E5F8"/>'
    '</w:tblBorders>'
    '</w:tblPr>' % NS
)

TC_MAR = (
    '<w:tcMar %s>'
    '<w:top w:w="80" w:type="dxa"/><w:start w:w="80" w:type="dxa"/>'
    '<w:bottom w:w="80" w:type="dxa"/><w:end w:w="80" w:type="dxa"/>'
    '</w:tcMar>' % NS
)

UNDERSCORE_RUN = re.compile(r'_{6,}')


def parse_inline(text):
    """Split markdown text into (text, bold) runs; strips ` and trims underscore blanks."""
    text = text.replace('`', '')
    text = UNDERSCORE_RUN.sub('', text)
    runs = []
    for i, part in enumerate(re.split(r'\*\*', text)):
        if part:
            runs.append((part, i % 2 == 1))
    return runs


def add_runs(paragraph, runs, italic=False, font=None, size=None):
    for text, bold in runs:
        run = paragraph.add_run(text)
        run.bold = bold or None
        run.italic = italic or None
        if font:
            run.font.name = font
        if size:
            run.font.size = size


def add_table(doc, rows, page_width_dxa):
    ncols = max(len(r) for r in rows)
    table = doc.add_table(rows=0, cols=ncols)
    tbl = table._tbl
    for existing in tbl.findall(qn('w:tblPr')):
        tbl.remove(existing)
    tbl.insert(0, parse_xml(TBL_PR))
    col_w = page_width_dxa // ncols
    for r_i, row in enumerate(rows):
        tr = table.add_row()
        for c_i in range(ncols):
            cell = tr.cells[c_i]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn('w:tcW'))
            if tc_w is None:
                tc_w = parse_xml('<w:tcW %s w:type="dxa" w:w="%d"/>' % (NS, col_w))
                tc_pr.append(tc_w)
            else:
                tc_w.set(qn('w:type'), 'dxa')
                tc_w.set(qn('w:w'), str(col_w))
            tc_pr.append(parse_xml('<w:vAlign %s w:val="top"/>' % NS))
            tc_pr.append(parse_xml(TC_MAR))
            if r_i == 0:
                tc_pr.append(parse_xml('<w:shd %s w:fill="EAF4FF"/>' % NS))
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            text = row[c_i] if c_i < len(row) else ''
            runs = parse_inline(text)
            if r_i == 0:
                runs = [(t, True) for t, _ in runs]
            add_runs(p, runs, font='Arial', size=Pt(9))


def split_table_row(line):
    return [c.strip() for c in line.strip().strip('|').split('|')]


def is_separator_row(cells):
    return all(re.fullmatch(r':?-{3,}:?', c) for c in cells if c)


def convert(md_path, docx_path, template_path, footer_text):
    shutil.copy(template_path, docx_path)
    doc = Document(docx_path)

    body = doc.element.body
    for child in list(body):
        if child.tag != qn('w:sectPr'):
            body.remove(child)

    # footer: keep formatting of the first run, replace the text
    footer = doc.sections[0].footer
    if footer.paragraphs:
        p = footer.paragraphs[0]
        if p.runs:
            p.runs[0].text = footer_text
            for extra in p.runs[1:]:
                extra.text = ''
        else:
            p.text = footer_text

    sect = doc.sections[0]
    usable = int((sect.page_width - sect.left_margin - sect.right_margin) / 635)  # EMU -> dxa

    lines = open(md_path).read().split('\n')
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if not stripped or re.fullmatch(r'-{3,}', stripped):
            i += 1
            continue

        m = re.match(r'^(#{1,4})\s+(.*)', stripped)
        if m:
            level = min(len(m.group(1)), 3)
            text = ' '.join(t for t, _ in parse_inline(m.group(2))).strip()
            doc.add_paragraph(text, style='Heading %d' % level)
            i += 1
            continue

        if stripped.startswith('|'):
            rows = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                cells = split_table_row(lines[i])
                if not is_separator_row(cells):
                    rows.append(cells)
                i += 1
            if rows:
                add_table(doc, rows, usable)
                doc.add_paragraph()
            continue

        if stripped.startswith('>'):
            quote_lines = []
            while i < len(lines) and lines[i].strip().startswith('>'):
                quote_lines.append(lines[i].strip().lstrip('>').strip())
                i += 1
            text = ' '.join(q for q in quote_lines if q).strip()
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Pt(18)
            if text:
                add_runs(p, parse_inline(text), italic=True, font='Arial')
            continue

        m = re.match(r'^[-*]\s+(.*)', stripped)
        if m:
            p = doc.add_paragraph(style='List Bullet')
            add_runs(p, parse_inline(m.group(1)))
            i += 1
            continue

        m = re.match(r'^(\d+)\.\s+(.*)', stripped)
        if m:
            p = doc.add_paragraph()
            add_runs(p, [('%s. ' % m.group(1), False)] + parse_inline(m.group(2)))
            i += 1
            continue

        p = doc.add_paragraph()
        add_runs(p, parse_inline(stripped))
        i += 1

    doc.save(docx_path)
    print('built', docx_path)


def main():
    jobs = eval(open(sys.argv[1]).read())
    for job in jobs:
        convert(job['md'], job['docx'], job['template'], job['footer'])


if __name__ == '__main__':
    main()
